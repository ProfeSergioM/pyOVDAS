from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
import datetime as dt
def REAV(a,pathmap,DR,hv=True):
    if hv==True:
        hvf=-3
    else:
        hvf=-4
    #prueba de hora local
    fecha=a.iloc[0]['ev_fecha']
    volcan_db=a.iloc[0]['nombredb'];reg=a.iloc[0]['vol_region'];tipoes=a.iloc[0]['vol_tipo'];volcan_re=a.iloc[0]['nombrevolcan']
    tipoev=a.iloc[0]['ev_tipoev'];lat_ev=a.iloc[0]['ev_lat'];lon_ev=a.iloc[0]['ev_lon'];prof_ev=a.iloc[0]['ev_prof'];ML=a.iloc[0]['ev_ml']
    ale=a.iloc[0]['vol_alerta']
    fecha = dt.datetime.strftime(fecha,'%Y-%m-%d %H:%M:%S')
    fecha_dt = dt.datetime.strptime(fecha, '%Y-%m-%d %H:%M:%S')
    dia = dt.datetime.now().strftime("%d")
    dia = int(dia)
    mes = dt.datetime.now().strftime("%m")
    anio = dt.datetime.now().strftime("%Y")
    hora = dt.datetime.now().strftime("%H")
    minuto =  dt.datetime.now().strftime("%M")
    mda = ["01","02","03","04","05","06","07","08","09","10","11","12" ]
    mda_str = ["enero","febrero","marzo","abril","mayo","junio","julio","agosto","septiembre","octubre","noviembre","diciembre" ]
    DICT_mes = dict(zip(mda,mda_str))
    wd,nwd = [0,1,2,3,4,5,6],["lunes","martes","miercoles","jueves","viernes",u"sábado","domingo"]
    DICT_wd = dict(zip(wd,nwd))
    d_dia = int(dt.datetime.now().strftime("%d")) - int(fecha[8:10])
    if d_dia==0:d_dia_str = u"Hoy"
    else:d_dia_str = u"El día" 
    hora_loc_hl = str(int(fecha_dt.strftime("%H"))+hvf)
    if hora_loc_hl =="-1":
        hora_loc_hl = "23"
        d_dia_str = u"El día" 
    elif hora_loc_hl =="-2":
        hora_loc_hl = "22"
        d_dia_str = u"El día" 
    elif hora_loc_hl =="-3":
        hora_loc_hl = "21"
        d_dia_str = u"El día" 
    elif hora_loc_hl =="-4":
        hora_loc_hl = "20"
        d_dia_str = u"El día" 
    #dia_loc = str(int(fecha_dt.strftime("%d")))
    #mes_loc = fecha_dt.strftime("%m")
    hora_loc = str(int(fecha_dt.strftime("%H")))
    #hora_loc_hl = str(int(fecha_dt.strftime("%H"))+hvf)
    
    
    min_loc_hl = str(int(fecha_dt.strftime("%M")))
    if len(hora_loc_hl) == 1:hora_loc_hl = "0"+hora_loc_hl
    if len(hora_loc) == 1:hora_loc = "0"+hora_loc
    if len(min_loc_hl) == 1: min_loc_hl = "0"+min_loc_hl

    document = Document('//172.16.40.10/Sismologia/pyOvdas_lib/templates/SNGM.docx')
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('gob', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'gobCL'
    obj_charstyle = obj_styles.add_style('gob2', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(13)
    obj_font.name = 'gobCL'
    obj_charstyle = obj_styles.add_style('gob3', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'gobCL'
	
    document._body.clear_content()
	#########################################################TITULO
    titulo1 = document.add_paragraph(u'')
    run = titulo1.add_run(u'Reporte Especial de Actividad Volcánica (REAV)', style='gob')
    run.bold = True
    titulo1_formato = titulo1.paragraph_format
    titulo1_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo1_formato.line_spacing = Pt(15)
    titulo1_formato.space_before = Pt(0)
    titulo1_formato.space_after = Pt(0)
    #########################################################TITULO 2 - REGION Y VOLCAN
    titulo2 = document.add_paragraph(u'')    
    run = titulo2.add_run(u"Región " + str(reg).title()+", "+tipoes+" "+volcan_re, style='gob2')
    titulo2_formato = titulo2.paragraph_format
    titulo2_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo2_formato.line_spacing = Pt(14)
    titulo2_formato.space_before = Pt(2)    
    #########################################################TITULO 2 - FECHA
    titulo2 = document.add_paragraph(u'')
    hoy_utc = dt.datetime.utcnow()    
    hoy_local = dt.datetime.utcnow() + dt.timedelta(hours=hvf)
    run = titulo2.add_run(hoy_local.strftime("%d")+" de "+DICT_mes[hoy_local.strftime("%m")]+ " de "+hoy_local.strftime("%Y")+", "+hoy_local.strftime("%H")+":"+hoy_local.strftime("%M")+" Hora local (Chile continental)", style='gob3')
    titulo2_formato = titulo2.paragraph_format
    titulo2_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.add_break(WD_BREAK.LINE_CLEAR_RIGHT)
    titulo2_formato.line_spacing = Pt(14)
    titulo2_formato.space_before = Pt(2)    
    #########################################################RESUMEN
    intro = document.add_paragraph(u'')
    #intro_formato = intro.paragraph_format
    run = intro.add_run(u"El ", style='gob3')
    run = intro.add_run(u"Servicio Nacional de Geología y Minería de Chile (Sernageomin) ", style='gob3')
    run.bold = True
    run = intro.add_run(u" da a conocer la siguiente información PRELIMINAR, obtenida a través de los equipos de monitoreo" +
						u" de la Red Nacional de Vigilancia Volcánica (RNVV)," +
						u" procesados y analizados en el Observatorio Volcanológico de los Andes del Sur (Ovdas):", style='gob3')
    intro_formato = intro.paragraph_format
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_formato.line_spacing = Pt(14)	    
	#########################################################intro evento
    resumen = document.add_paragraph(u'')
    run = resumen.add_run(d_dia_str+" "+DICT_wd[fecha_dt.weekday()] +" "+fecha[8:10]+" de "+DICT_mes[fecha[5:7]]+" a las "
                          +hora_loc_hl+":"+min_loc_hl+" hora local ("+hora_loc+":"+min_loc_hl+
						u" UTC), las estaciones de monitoreo instaladas en las inmediaciones del ", style='gob3')
    run = resumen.add_run(tipoes+" "	+volcan_re, style='gob3')
    run.bold = True    
 
    if tipoev == "VT":tte= u" registraron un sismo asociado al fracturamiento de roca (Volcano-Tectónico). "
    elif tipoev == "LP":tte= u" registraron un sismo asociado a la dinámica de fluidos al interior del sistema volcánico (Largo Periodo). "
    elif tipoev == "HB":tte= u" registraron un sismo asociado tanto al fracturamiento de roca como a la dinámica de fluidos al interior del sistema volcánico (Híbrido). "
    elif tipoev == "EX":tte= u" registraron una explosión asociada a la dinámica de fluidos al interior del sistema volcánico con emisión de gases y de material particulado. "
    elif tipoev == "VD":tte= u" registraron un sismo lejano. "
    run = resumen.add_run(tte, style='gob3')					
    resumen_formato = resumen.paragraph_format
    resumen.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    resumen_formato.line_spacing = Pt(14)
    resumen2 = document.add_paragraph(u'')
    run = resumen2.add_run(u"Las características del sismo luego de su análisis son las siguientes:", style='gob3')
	#########################################################MAPA
    locima = document.add_paragraph(u'')
    run = locima.add_run()
    run.add_picture(pathmap,width=Inches(5.2))
    locima.paragraph_format
    locima.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #########################################################tabla datos
    datos = document.add_paragraph(u'')
    run = datos.add_run(u"TIEMPO DE ORIGEN:\t\t\t")
    run = datos.add_run(hora_loc_hl+":"+min_loc_hl+" hora local ("+hora_loc+":"+min_loc_hl+ u" UTC)")
    run.add_break(WD_BREAK.LINE)
    run = datos.add_run(u"LATITUD:\t\t\t\t")
    run = datos.add_run(str(lat_ev).replace(".",",")[1:7]+u"\xb0 S")  
    run.add_break(WD_BREAK.LINE)
    run = datos.add_run(u"LONGITUD:\t\t\t\t")
    run = datos.add_run(str(lon_ev).replace(".",",")[1:7]+u"\xb0 O")    
    run.add_break(WD_BREAK.LINE)
    run = datos.add_run(u"PROFUNDIDAD:\t\t\t\t")
    run = datos.add_run(str(prof_ev).replace(".",",")[:-1]+" km") 
    run.add_break(WD_BREAK.LINE)
    if tipoev in ("VT","VD"):
        run = datos.add_run(u"MAGNITUD LOCAL:\t\t\t")
        run = datos.add_run(str(ML).replace(".",",")+" (ML)") 
    elif tipoev == "LP":
        run = datos.add_run(u"DESPLAZAMIENTO REDUCIDO:\t\t")
        run = datos.add_run(str(DR).replace(".",",")+" (cm*cm)") 
    datos_formato = datos.paragraph_format
    datos_formato.line_spacing = Pt(14)
    run.add_break(WD_BREAK.LINE)
    obs = document.add_paragraph(u'')
    run = obs.add_run(u"OBSERVACIONES:", style='gob3')
    run.bold = True
    run.add_break(WD_BREAK.LINE)
    run.add_break(WD_BREAK.LINE)
    run = obs.add_run(u"INSERTE AQUÍ LAS OBSERVACIONES - INSERTE AQUÍ LAS OBSERVACIONES - " +
					u"INSERTE AQUÍ LAS OBSERVACIONES - INSERTE AQUÍ LAS OBSERVACIONES - " +
					u"INSERTE AQUÍ LAS OBSERVACIONES - INSERTE AQUÍ LAS OBSERVACIONES - " +
					u"INSERTE AQUÍ LAS OBSERVACIONES - INSERTE AQUÍ LAS OBSERVACIONES - " +
					u"INSERTE AQUÍ LAS OBSERVACIONES - INSERTE AQUÍ LAS OBSERVACIONES - " +
					u"INSERTE AQUÍ LAS OBSERVACIONES - INSERTE AQUÍ LAS OBSERVACIONES - ", style='gob3')
    obs_formato = obs.paragraph_format
    obs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    obs_formato.line_spacing = Pt(14)
    run.add_break(WD_BREAK.LINE)
    alerta = document.add_paragraph(u'')
    if ale==3:
        rutale='NN.png'
    elif ale==4:
        rutale='NR.png'
    elif ale==2:
        rutale='NA.png'
    elif ale==1:
        rutale='NV.png'
    run = alerta.add_run(u"La alerta técnica volcánica se mantiene en:", style='gob3')    
	#########################################################MAPA
    locima = document.add_paragraph(u'')
    run = locima.add_run()
    run.add_picture('//172.16.40.10/Sismologia/pyOvdas_lib/templates/'+rutale,width=Inches(3))
    locima.paragraph_format
    locima.alignment = WD_ALIGN_PARAGRAPH.CENTER
    alerta2 = document.add_paragraph(u'')
    run = alerta2.add_run(u"Sernageomin realiza vigilancia en linea e informa de manera oportuna" 
						u" sobre eventuales cambios en la actividad volcánica del país.", style='gob3')
    alerta2.paragraph_format
    alerta2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	
    pie = document.add_paragraph(u'')
	
    run = pie.add_run(u"Servicio Nacional de Geología y Minería (Sernageomin)", style='gob3')
    run.bold = True
    run.add_break(WD_BREAK.LINE)
    run = pie.add_run(u"Red Nacional de Vigilancia Volcánica (RNVV)", style='gob3')
    run.add_break(WD_BREAK.LINE)
    run = pie.add_run(u"Observatorio Volcanológico De los Andes del Sur (Ovdas)", style='gob3')
    run.add_break(WD_BREAK.LINE)
    pie_formato = pie.paragraph_format
    pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pie_formato.line_spacing = Pt(14)    
    
    
    
    path='REAV_'+fecha[0:4]+fecha[5:7]+fecha[8:10]+"_"+hoy_local.strftime("%H")+hoy_local.strftime("%M")+"_"+volcan_db+'.docx'

    document.save(path)